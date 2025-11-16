#!/usr/bin/env python3
"""
Web Interface for Booking Confirmation Transformer
Drag-and-drop PDFs, Word docs, or paste text to generate confirmations
"""

import os
import json
import tempfile
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_file, session
from werkzeug.utils import secure_filename
from claude_transform import CloudbedsTransformer
import secrets

app = Flask(__name__, template_folder='web_templates')
app.secret_key = secrets.token_hex(16)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = Path('temp_uploads')
app.config['UPLOAD_FOLDER'].mkdir(exist_ok=True)

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    """Main upload page"""
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def upload():
    """Handle file upload or text paste"""
    try:
        # Check if file or text was provided
        if 'file' in request.files and request.files['file'].filename:
            file = request.files['file']

            if not allowed_file(file.filename):
                return jsonify({'error': 'Invalid file type. Allowed: PDF, DOCX, DOC, TXT'}), 400

            # Save file temporarily
            filename = secure_filename(file.filename)
            filepath = app.config['UPLOAD_FOLDER'] / filename
            file.save(filepath)

            # Extract data using Claude API
            transformer = CloudbedsTransformer(debug=True)

            if filename.lower().endswith('.pdf'):
                booking_data = transformer.extract_booking_data(str(filepath))
            elif filename.lower().endswith(('.docx', '.doc')):
                # Extract text from Word doc (paragraphs AND tables)
                from docx import Document
                doc = Document(filepath)

                # Extract from paragraphs
                text_parts = [para.text for para in doc.paragraphs if para.text.strip()]

                # Extract from tables (where most booking data usually is)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            text_parts.append(row_text)

                text_content = '\n'.join(text_parts)

                # Validate extraction
                print(f"📄 Extracted {len(text_content)} characters from Word document")
                print(f"📝 Preview: {text_content[:200]}...")

                if len(text_content) < 50:
                    raise ValueError(f"Word document appears empty (only {len(text_content)} characters extracted). Please check the file format.")

                booking_data = transformer.extract_from_text(text_content)
            else:  # .txt
                with open(filepath, 'r', encoding='utf-8') as f:
                    text_content = f.read()

                print(f"📄 Text file: {len(text_content)} characters")

                if len(text_content) < 50:
                    raise ValueError(f"Text file appears empty (only {len(text_content)} characters). Please check the file content.")

                booking_data = transformer.extract_from_text(text_content)

            # Clean up temp file
            filepath.unlink()

        elif 'text_content' in request.form and request.form['text_content'].strip():
            # Handle pasted text
            text_content = request.form['text_content'].strip()

            print(f"📝 Pasted text: {len(text_content)} characters")
            print(f"📝 Preview: {text_content[:200]}...")

            if len(text_content) < 50:
                return jsonify({'error': f'Text is too short ({len(text_content)} characters). Please paste the complete booking confirmation.'}), 400

            transformer = CloudbedsTransformer(debug=True)
            booking_data = transformer.extract_from_text(text_content)

            # Validate extraction results
            if not booking_data or not booking_data.get('guest_name'):
                return jsonify({'error': 'Could not extract booking data. Please check the text format and try again.'}), 400
        else:
            return jsonify({'error': 'Please upload a file or paste text'}), 400

        # Validate booking data has minimum required fields
        # Email is NOT required for agent bookings (guest hasn't arrived yet)
        if booking_data.get('booking_type') == 'agent':
            required_fields = ['guest_name', 'check_in', 'check_out', 'res_id']
        else:
            required_fields = ['guest_name', 'email', 'check_in', 'check_out', 'res_id']

        missing_fields = [f for f in required_fields if not booking_data.get(f)]

        if missing_fields:
            print(f"⚠️ Warning: Missing fields: {', '.join(missing_fields)}")
            # Don't fail, but log the warning

        print(f"✅ Extraction successful: {booking_data.get('guest_name', 'Unknown guest')}")

        # Store booking data in session for review
        session['booking_data'] = booking_data

        return jsonify({
            'success': True,
            'booking_data': booking_data,
            'redirect': '/review'
        })

    except ValueError as e:
        # Validation errors (empty files, etc.)
        print(f"❌ Validation error: {str(e)}")
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        # Unexpected errors
        print(f"❌ Unexpected error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Extraction failed: {str(e)}'}), 500

@app.route('/review')
def review():
    """Review extracted booking data"""
    booking_data = session.get('booking_data')
    if not booking_data:
        return "No booking data found. Please upload a file first.", 400

    return render_template('review.html', data=booking_data)

@app.route('/update', methods=['POST'])
def update():
    """Update booking data from review form"""
    try:
        updated_data = request.json
        session['booking_data'] = updated_data
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generate', methods=['POST'])
def generate():
    """Generate final confirmation"""
    try:
        booking_data = session.get('booking_data')
        if not booking_data:
            return jsonify({'error': 'No booking data found'}), 400

        # Generate confirmation
        transformer = CloudbedsTransformer(debug=True)
        html_output, pdf_output = transformer.generate_confirmation(booking_data)

        # Store output paths in session
        session['html_output'] = str(html_output)
        session['pdf_output'] = str(pdf_output)

        return jsonify({
            'success': True,
            'html_file': html_output.name,
            'pdf_file': pdf_output.name if pdf_output else None,
            'confirmation_number': booking_data.get('res_id', 'N/A')
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download/<file_type>')
def download(file_type):
    """Download generated HTML or PDF"""
    try:
        if file_type == 'html':
            filepath = session.get('html_output')
        elif file_type == 'pdf':
            filepath = session.get('pdf_output')
        else:
            return "Invalid file type", 400

        if not filepath or not Path(filepath).exists():
            return "File not found", 404

        return send_file(filepath, as_attachment=True)

    except Exception as e:
        return str(e), 500

@app.route('/success')
def success():
    """Success page with download links"""
    html_file = session.get('html_output')
    pdf_file = session.get('pdf_output')

    if not html_file:
        return "No confirmation generated yet.", 400

    return render_template('success.html',
                         html_file=Path(html_file).name,
                         pdf_file=Path(pdf_file).name if pdf_file else None)

if __name__ == '__main__':
    print("=" * 60)
    print("🦚 The Planters House - Booking Confirmation Generator")
    print("=" * 60)
    print("\n📱 Web Interface Starting...")
    print("🌐 Open: http://localhost:5000")
    print("📝 Drag & drop PDFs, Word docs, or paste text")
    print("\n✅ Ready to transform booking confirmations!\n")

    app.run(debug=True, host='0.0.0.0', port=5000)
